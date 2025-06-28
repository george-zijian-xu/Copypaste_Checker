#!/usr/bin/env python3
"""
docx_to_text.py

A utility to convert a .docx file into a plain text file.
It extracts text from both paragraphs and tables.

Requires the 'python-docx' library.
Install it using: pip install python-docx
"""
import sys
import os
try:
    import docx
except ImportError:
    print("Error: The 'python-docx' library is required.")
    print("Please install it using: pip install python-docx")
    sys.exit(1)

def convert_docx_to_text(docx_path, txt_path):
    """
    Extracts all text from a .docx file and saves it to a .txt file.
    Includes text from paragraphs and tables.
    """
    if not os.path.isfile(docx_path):
        print(f"Error: Input file not found at '{docx_path}'")
        return False
    
    try:
        document = docx.Document(docx_path)
        full_text = []

        # Extract text from paragraphs
        for para in document.paragraphs:
            full_text.append(para.text)

        # Extract text from tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write('\n'.join(full_text))
            
        return True

    except Exception as e:
        print(f"An unexpected error occurred during conversion: {e}")
        return False

def main():
    if len(sys.argv) != 3:
        print("Usage: python docx_to_text.py <input.docx> <output.txt>")
        sys.exit(1)
        
    docx_path, txt_path = sys.argv[1:3]
    
    print(f"Converting '{docx_path}' to plain text...")
    if convert_docx_to_text(docx_path, txt_path):
        print(f"Successfully converted document to '{txt_path}'")
    else:
        print("Conversion failed.")

if __name__ == "__main__":
    main()